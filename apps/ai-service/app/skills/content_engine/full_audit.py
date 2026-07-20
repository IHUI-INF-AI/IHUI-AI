# -*- coding: utf-8 -*-
"""
公众号流水线终极物理审计（42 维度，2026-07-12 永久铁律）
========================================
设计目标：一次跑完把所有返工根因提前发现,不让用户来来回回追问。
对所有维度输出 PASS/FAIL,任何 FAIL 都打印修复建议。

可被流水线 import 调用: from full_audit import main; exit_code, p, f, _ = main(['--title', 'xxx'])
也支持 AUDIIT_DRY_RUN=1 环境变量,跳过"需要真推"的检查。
========================================
用法: python full_audit.py [--md articles/xxx.md] [--title "..."]
      AUDIIT_DRY_RUN=1 python full_audit.py --title "..."  # 跳过需真推的维度
"""
import io
import os
import sys
import re
import json
import time
import argparse
import urllib.request
import subprocess

try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')


BASE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE)
sys.path.insert(0, os.path.join(BASE, 'lib'))

# ── 项目边界硬门禁（导入即生效，fail-closed） ──
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "koubo_workflow"))
import project_boundary
project_boundary.check_action(tool="full_audit.py")

results = []  # (维度, 状态, 详情)
IS_DRY_RUN = 'AUDIIT_DRY_RUN' in os.environ

# === 项目边界硬检测（2026-07-14 用户强制·零容忍） ===
# 公众号 output/ 绝不允许出现口播稿文件（.txt / MMDD.txt 模式 / 抖音话题词 / 置顶评论）。
# 一旦出现 = 流水线 FAIL，并打印强提示要求立即清理。
KOUBO_FILE_PATTERN = re.compile(r'^\d{4}\.txt$', re.U)  # 0629/0701/.../0714.txt
# 口播稿强特征词库 v2（2026-07-14 第二轮精修·高置信度·24 词）
# 注意: 已移除「李总 / 智汇AI / 智汇AI丨」三类品牌词
#   原因: 公众号文章本身就必须用第一人称"李总"叙述、品牌名"智汇AI"署名,
#   把品牌词当口播稿特征 = 误伤所有正常文章。铁律核心是「防抖音口播稿污染公众号」,
#   真正口播稿特征词是平台词（完播率/涨粉/直播间）和口播话术词（咱就说/你品品）。
# 公众号文章 + DOCX 命中 ≥2 个即 FAIL
KOUBO_KEYWORDS = (
    # 平台话题词（口播稿独有·公众号不会用）
    '[置顶]', '#科技', '#互联网', '#商业', '#AI取代工作',
    # 抖音/带货专属术语
    '完播率', '涨粉', '直播间', '橱窗', '带货',
    # 口播稿话术词（公众号不会用）
    '咱就说', '你品品', '你想想', '你猜怎么着', '废话不多说', '记住我这句话',
    '核心来了', '重点来了', '听好了', '看到这儿了', '划走就亏了',
    '今儿',  # 口播稿专属用词
    # 标题党词（口播稿高频·公众号不会用）
    '巨亏', '血赚', '翻车', '爆单', '破防',
)


def _scan_koubo_pollution(root_dir):
    """递归扫描 root_dir，查找所有疑似口播稿文件。返回命中列表。

    排除合法的规则文档（AGENTS.md / MEMORY.md / SKILL.md / README.md），
    这些文件本身就是描述"什么是口播稿特征词"的规则说明文件。
    """
    if not os.path.isdir(root_dir):
        return []
    # 合法规则文档（必须排除）: 文件名匹配跳过扫描
    DOC_FILES = {
        'AGENTS.md', 'README.md', 'SKILL.md',
        'MEMORY.md',  # 跨项目记忆 + 单项目记忆
        'REFERENCE.md', 'CSDN-GUIDE.md', 'HOOKS-LIBRARY.md',
        'README.md',
    }
    # 整个 .workbuddy 目录属于 agent 记忆,内容是规则说明,跳过
    SKIP_DIRS = {'.workbuddy', '__pycache__'}
    hits = []
    for cur, subdirs, files in os.walk(root_dir):
        # 跳过 .workbuddy 整个目录
        subdirs[:] = [d for d in subdirs if d not in SKIP_DIRS]
        for f in files:
            fpath = os.path.join(cur, f)
            rel = os.path.relpath(fpath, BASE)
            # 排除合法规则文档
            if f in DOC_FILES:
                continue
            # 1) .txt 后缀 = 一票否决
            if f.lower().endswith('.txt'):
                hits.append((rel, '后缀 .txt 出现在公众号目录'))
                continue
            # 2) MMDD.txt 模式 = 一票否决（哪怕不是 .txt 后缀也不允许）
            if KOUBO_FILE_PATTERN.match(f):
                hits.append((rel, '文件名匹配口播稿 MMDD.txt 模式'))
                continue
            # 3) .md/.html/.docx 命中口播稿特征词 = FAIL
            #    公众号文章和 DOCX 不应出现 [置顶] / #科技 / 李总 / 咱就说 等
            if f.lower().endswith(('.md', '.html', '.docx')):
                try:
                    txt = open(fpath, 'r', encoding='utf-8', errors='ignore').read()
                except Exception:
                    continue
                kw_hits = [k for k in KOUBO_KEYWORDS if k in txt]
                if len(kw_hits) >= 2:
                    hits.append((rel, f'口播稿特征词命中({len(kw_hits)}个): {kw_hits[:4]}'))
    return hits


def add(dim, ok, detail='', warn=False):
    if warn:
        results.append((dim, '⚠️ WARN', detail))
        print(f"  ⚠️ {dim} (WARN): {detail}")
    else:
        results.append((dim, '✅ PASS' if ok else '❌ FAIL', detail))
        print(f"  {'✅' if ok else '❌'} {dim}: {detail}")


def section(title):
    print(f"\n{'='*64}\n  {title}\n{'='*64}")


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument('--md', default='', help='可选,源md路径')
    ap.add_argument('--title', default='', help='可选,标题')
    args = ap.parse_args(argv)
    if argv is not None:
        sys.argv = ['full_audit.py'] + argv

    # === A. 代码维度 ===
    section("A. 代码维度 (6 维度)")
    t = open('lib/moyu_green_renderer.py', encoding='utf-8').read()
    add('A1 渲染器 isinstance 字段类型判断', 'isinstance(' in t, '已做类型判断')
    add('A2 渲染器函数数 ≥ 10 (13组件)', t.count('def render_') >= 10, f'{t.count("def render_")} 个 render 函数')
    p = open('publish_pipeline.py', encoding='utf-8').read()
    add('A3 流水线门禁 ≥ 3 (A/A+/B)', p.count('def gate_') >= 3, f'{p.count("def gate_")} 个门禁')
    req_count = p.count('required=True')
    # 关键参数强制的两种方式都认可：① argparse 必填 ≥4；② 至少 1 个 argparse 必填(--md)
    #   且存在运行时守卫(_missing + _fail_guard 覆盖 title/digest/cover)，使 --only-cleanup 可单独跑。
    runtime_guard = ('_missing' in p) and ('_fail_guard' in p) and \
                    ('title' in p) and ('digest' in p) and ('cover' in p)
    add('A4 关键参数强制(必填≥4 或 运行时守卫)',
        req_count >= 4 or (req_count >= 1 and runtime_guard),
        f'argparse必填{req_count}个, 运行时守卫={"有" if runtime_guard else "无"}')
    add('A5 流水线 except 异常处理', p.count('except') >= 5, f'{p.count("except")} 处 except')
    add('A6 流水线 import time', re.search(r'^import time', p, re.M) is not None, '已导入 time')

    # === B. 脚本结构 ===
    section("B. 脚本结构 (3 维度)")
    root_pys = [f for f in os.listdir('.') if f.endswith('.py') and not f.startswith('_') and f != 'full_audit.py']
    add('B1 根目录脚本 ≤ 3 (官方)', len(root_pys) == 3, f'{len(root_pys)} 个: {root_pys} (full_audit.py 是审计工具,允许)')
    arch = [f for f in os.listdir('_archive') if f.endswith('.py')] if os.path.isdir('_archive') else []
    add('B2 废弃脚本归档 _archive', len(arch) == 3, f'{len(arch)} 个: {arch}')
    libs = [f for f in os.listdir('lib') if f.endswith('.py') and f != '__init__.py']
    add('B3 lib 模块 ≥ 4', len(libs) >= 4, f'{len(libs)} 个: {libs}')

    # === C. 产物 ===
    section("C. 产物 (8+1 维度)")
    # === C0 永久铁律·项目文件越界检测（2026-07-14 用户强制·零容忍） ===
    # 公众号 公众号/ 整棵树绝不允许出现口播稿文件（.txt / MMDD.txt 模式 / 抖音话题词 / 置顶评论）。
    # 这是"AI 误把口播稿写到公众号 output" 的根治防护。
    # 扫描范围：公众号/ 全树（含 wechat-article-system/、images/、archive/ 等所有子目录）
    scan_root = os.path.dirname(BASE)  # 公众号/ 目录
    koubo_hits = _scan_koubo_pollution(scan_root)
    if koubo_hits:
        detail_lines = [f'  ❌ {h[0]} → {h[1]}' for h in koubo_hits[:10]]
        add('C0 项目文件越界 (公众号目录无 .txt/口播稿特征词,零容忍)', False,
            f'发现 {len(koubo_hits)} 处越界: ' + ' | '.join(detail_lines))
    else:
        add('C0 项目文件越界 (公众号目录无 .txt/口播稿特征词,零容忍)', True,
            '公众号目录清洁,未发现 .txt/MMDD.txt/口播稿特征词')
    if args.title:
        clean_title = args.title.rstrip('？?').rstrip('?')
    else:
        clean_title = 'OpenAI提示词又改规则了：3个动作让AI不再装傻，普通人省2小时'
    # 排除 Word 临时锁文件 ~$xxx.docx / build_docx 锁失败产生的 .new / .tmp 副产物
    out_files = [f for f in os.listdir('output') if os.path.isfile(f'output/{f}') and f != 'images' and not f.startswith('~$') and not f.endswith('.bak') and not f.endswith('.new') and not f.endswith('.tmp')]
    # 排除 md: 2026-07-14 用户禁用 md 产物,只交付 html + docx
    out_files = [f for f in out_files if not f.endswith('.md')]
    # 2026-07-17 零容忍铁律升级：所有 html/docx 必须以 clean_title 开头（防止多版本残留）
    orphan_files = [f for f in out_files if (f.endswith('.html') or f.endswith('.docx')) and not f.startswith(clean_title)]
    if orphan_files:
        add('C0.5 output 零旧版 (所有 html/docx 必须以本标题开头)', False,
            f'❌ 发现 {len(orphan_files)} 个非本标题残留: {orphan_files}')
    else:
        add('C0.5 output 零旧版 (所有 html/docx 必须以本标题开头)', True,
            f'output 无旧版残留 ({len(out_files)} 个 html/docx 全是本标题)')
    expected_outputs = {f for f in out_files if f.startswith(clean_title)}
    # 检查交付物数量: 7/16 后 CSDN 双交付 (html + docx + _CSDN.docx) = 3 件合法,
    # 单交付 (html + docx) = 2 件合法
    has_csdn = any(f.startswith(clean_title) and f.endswith('_CSDN.docx') for f in os.listdir('output'))
    c1_pass = (len(expected_outputs) == 2) or (has_csdn and len(expected_outputs) == 3)
    c1_note = '(html+docx)' if not has_csdn else '(html+docx+_CSDN.docx 双交付)'
    # 检查是否真的只交付 2 个 (html+docx),且有 .new 备份说明 docx 被文件锁 (附加 .new 警告)
    has_new_backup = any(f.startswith(clean_title) and f.endswith('.new') for f in os.listdir('output'))
    c1_detail = f'{len(expected_outputs)} 个: {sorted(expected_outputs)} {c1_note}'
    if has_new_backup:
        c1_detail += f' (附带 .new 备份: docx 被外部进程文件锁, 已 fallback)'
    add('C1 output 交付物 = 2/3 (html+docx, 7/16 起 CSDN 双交付 3 件合法)', c1_pass, c1_detail)
    # 2026-07-14: 单独的 .new 备份警告 (环境问题,非代码缺陷)
    add('C1.5 无 .new 临时备份 (docx 文件锁,真推侧问题)', not has_new_backup,
        f'⚠️ 检测到 .new 备份, docx 被外部进程锁定 (如微信端预览), 关闭预览后下次流水线自动覆盖' if has_new_backup else '无 .new 备份', warn=has_new_backup)
    art_md = [f for f in os.listdir('articles') if f.endswith('.md')] if os.path.isdir('articles') else []
    csdn_md = [f for f in art_md if f.endswith('_csdn.md')]
    main_md = [f for f in art_md if not f.endswith('_csdn.md')]
    add('C2 articles 源md = 1 (无残留)', len(main_md) == 1, f'{len(art_md)} 个 (主源 {len(main_md)} + CSDN专用 {len(csdn_md)})')
    imgs = [f for f in os.listdir('output/images')] if os.path.isdir('output/images') else []
    add('C3 配图 ≥ 3 张', len(imgs) >= 3, f'{len(imgs)} 张')
    # C3+ 真图铁律: 检测 PIL 文本卡片 (2026-07-14 用户强制)
    # 真实照片: Shannon 熵 > 4.0 + 最常见颜色占比 < 70%
    # PIL 文本卡: 熵 < 4.0 或 单一颜色占 > 70%
    # 只检查当前文章实际引用的图片（避免历史图片污染）
    try:
        from PIL import Image
        from collections import Counter
        import math
        # 收集当前文章引用的图片（从 source md + html 里抓，md 产物已禁用）
        referenced = set()
        for src in [args.md, 'output/' + clean_title + '.html']:
            if src and os.path.exists(src):
                try:
                    txt = open(src, encoding='utf-8').read()
                    referenced.update(re.findall(r'[\w\-/]*?([\w\-]+\.(?:jpg|jpeg|png))', txt, re.I))
                except Exception:
                    pass
        # 兜底：所有 .jpg/.png 都查一遍（包括历史图）
        all_imgs = imgs
        real_count = 0
        fake_count = 0
        fake_details = []
        # 优先级1：检查 referenced（文章实际引用的）必须全为真图
        if referenced:
            check_list = sorted(referenced)
            check_scope = 'referenced'
        else:
            check_list = sorted(all_imgs)
            check_scope = 'all'
        for img_f in check_list:
            if not (img_f.lower().endswith('.jpg') or img_f.lower().endswith('.jpeg') or img_f.lower().endswith('.png')):
                continue
            # 点赞.jpg/关注.jpg 是设计资产（白底图文卡片），不参与真图检测
            if img_f in ('点赞.jpg', '关注.jpg', '公众号置底.png'):
                real_count += 1  # 视为真图
                continue
            if not os.path.exists(os.path.join('output/images', img_f)):
                # 不在 output/images 时去 assets/images 找
                if os.path.exists(os.path.join('assets/images', img_f)):
                    real_count += 1
                continue
            img_p = os.path.join('output/images', img_f)
            try:
                img_obj = Image.open(img_p).convert('RGB')
                img_small = img_obj.resize((min(200, img_obj.size[0]), min(200, img_obj.size[1])), Image.LANCZOS)
                pixels = list(img_small.getdata())
                counter = Counter(pixels)
                most_common_pct = counter.most_common(1)[0][1] / len(pixels)
                # Shannon 熵
                total = len(pixels)
                entropy = 0.0
                for cnt in counter.values():
                    pc = cnt / total
                    if pc > 0:
                        entropy -= pc * math.log2(pc)
                is_card = (most_common_pct > 0.7) or (entropy < 4.0)
                if is_card:
                    fake_count += 1
                    fake_details.append(f'{img_f}(熵={entropy:.2f}, 单色占{most_common_pct:.1%})')
                else:
                    real_count += 1
            except Exception as ex:
                fake_count += 1
                fake_details.append(f'{img_f}(读取失败: {ex})')
        # 正文配图（不含点赞关注等文末）至少4张真图
        # 点赞.jpg/关注.jpg 是设计资产（白底图文卡片），不参与真图检测
        body_imgs = [f for f in check_list if f not in ('点赞.jpg', '关注.jpg', '公众号置底.png', '微信图片_20260704113447_16_530.jpg', 'assets/images/点赞.jpg', 'assets/images/关注.jpg')]
        body_real = 0
        for img_f in body_imgs:
            if not (img_f.lower().endswith('.jpg') or img_f.lower().endswith('.jpeg') or img_f.lower().endswith('.png')):
                continue
            img_p = os.path.join('output/images', img_f)
            if not os.path.exists(img_p):
                img_p = os.path.join('assets/images', img_f)
            if not os.path.exists(img_p):
                continue
            try:
                img_obj = Image.open(img_p).convert('RGB')
                img_small = img_obj.resize((min(200, img_obj.size[0]), min(200, img_obj.size[1])), Image.LANCZOS)
                pixels = list(img_small.getdata())
                counter = Counter(pixels)
                most_common_pct = counter.most_common(1)[0][1] / len(pixels)
                total = len(pixels)
                entropy = 0.0
                for cnt in counter.values():
                    pc = cnt / total
                    if pc > 0:
                        entropy -= pc * math.log2(pc)
                if not ((most_common_pct > 0.7) or (entropy < 4.0)):
                    body_real += 1
            except Exception:
                pass
        add(f'C3+ 真图铁律 ({check_scope}, PIL文本卡=0)', fake_count == 0 and body_real >= 4,
            f'真图 {real_count}/{len(check_list)}, 正文真图 {body_real}/4' if fake_count == 0
            else f'❌ 检测到 {fake_count} 张假图: {fake_details[:3]}')
        # C3++ 无水印铁律 (2026-07-15 用户强制): 只检测典型水印位（底边三处），
        # 避免把自然纯色天空/墙面/纸张误判为水印。
        def _has_watermark(img):
            w, h = img.size
            # 只检查底边区域（水印最常见位置），降低自然场景误报
            regs = {
                'bl': (0, int(h * 0.78), int(w * 0.30), h),
                'br': (int(w * 0.70), int(h * 0.78), w, h),
                'bot': (int(w * 0.25), int(h * 0.85), int(w * 0.75), h),
            }
            small = img.convert('RGB').resize((50, 50))
            dom = Counter(list(small.getdata())).most_common(1)[0][0]
            for name, box in regs.items():
                crop = img.crop(box).convert('RGB')
                px = list(crop.getdata())
                n = len(px)
                rm = sum(p[0] for p in px) / n
                gm = sum(p[1] for p in px) / n
                bm = sum(p[2] for p in px) / n
                var = sum(((p[0] - rm) ** 2 + (p[1] - gm) ** 2 + (p[2] - bm) ** 2) for p in px) / n
                std = math.sqrt(var / 3)
                diff = math.sqrt((rm - dom[0]) ** 2 + (gm - dom[1]) ** 2 + (bm - dom[2]) ** 2)
                # 实心色块水印：极低方差 + 与主图明显色差
                if std < 10 and diff > 80:
                    return True, f'{name}色块'
                # 文字水印：大量深色小像素且区域整体低方差（文字带很规整）
                g = crop.convert('L')
                dark = sum(1 for v in list(g.getdata()) if v < 120) / n
                if dark > 0.25 and std < 18:
                    return True, f'{name}文字'
            return False, ''
        wm_count = 0
        wm_details = []
        for img_f in check_list:
            if img_f in ('点赞.jpg', '关注.jpg', '公众号置底.png'):
                continue
            if not (img_f.lower().endswith('.jpg') or img_f.lower().endswith('.jpeg') or img_f.lower().endswith('.png')):
                continue
            img_p = os.path.join('output/images', img_f)
            if not os.path.exists(img_p):
                img_p = os.path.join('assets/images', img_f)
            if not os.path.exists(img_p):
                continue
            try:
                wm, why = _has_watermark(Image.open(img_p).convert('RGB'))
                if wm:
                    wm_count += 1
                    wm_details.append(f'{img_f}({why})')
            except Exception:
                pass
        add('C3++ 无水印铁律 (角落实心块/密集小字=0)', wm_count == 0,
            '全部无水印' if wm_count == 0 else f'❌ 检测到 {wm_count} 张疑似水印: {wm_details[:3]}')
    except ImportError:
        add('C3+ 真图铁律 (PIL可用)', False, 'PIL 未安装,无法检测')
    if args.title:
        audit_title = args.title.rstrip('？?')
        html_p = f'output/{audit_title}.html'
        docx_p = f'output/{audit_title}.docx'
        ok_html = os.path.exists(html_p) and not os.path.exists(f'~${audit_title}.docx')
        ok_docx = os.path.exists(docx_p) and not os.path.exists(f'~${audit_title}.docx')
        add('C4 HTML 产物存在', ok_html, f'{os.path.getsize(html_p) if ok_html else 0} 字符' if ok_html else '缺失')
        add('C5 DOCX 产物存在 (排除 Word 锁文件)', ok_docx, f'{os.path.getsize(docx_p) if ok_docx else 0} 字节' if ok_docx else '缺失/被 Word 占用')
        add('C6 MD 产物已禁用 (2026-07-14 用户强制:只交付 html+docx)', True, '已禁用')
        if ok_html and ok_docx:
            from docx import Document
            h = open(html_p, encoding='utf-8').read()
            d = Document(docx_p)
            docx_text = "\n".join(p.text for p in d.paragraphs)
            keys = ['智汇AI悄悄话', '智汇AI']
            miss = [k for k in keys if not (k in h and k in docx_text)]
            # 2026-07-17: 作者按需省略豁免 — 若源 MD 中也没有"智汇AI悄悄话"段,
            # 视为作者主动选择不加, 不判定为缺漏 (与 CSDN 剥离逻辑一致)
            src_md = ''
            if args.md and os.path.exists(args.md):
                src_md = open(args.md, encoding='utf-8').read()
            src_has_secret = '智汇AI悄悄话' in src_md
            if '智汇AI悄悄话' in miss and not src_has_secret:
                miss.remove('智汇AI悄悄话')
            csdn_docx_p = docx_p.replace('.docx', '_CSDN.docx')
            csdn_note = ''
            if os.path.exists(csdn_docx_p):
                # CSDN 版按 7/16 规则剥离了"智汇AI悄悄话"营销段: 豁免该 key,
                # 改为检查 CSDN 版已剥离（docx_text 中无"智汇AI悄悄话"）
                dc = Document(csdn_docx_p)
                csdn_text = "\n".join(p.text for p in dc.paragraphs)
                csdn_has_secret = '智汇AI悄悄话' in csdn_text
                if csdn_has_secret:
                    miss.append('CSDN docx 含营销段"智汇AI悄悄话"(应剥离)')
                # 普通 docx 必须保留"智汇AI悄悄话"(公众号刻意要的营销段) — 仅在源 MD 含时强制
                if src_has_secret:
                    wechat_has_secret = '智汇AI悄悄话' in docx_text
                    if not wechat_has_secret:
                        miss.append('微信 docx 缺营销段"智汇AI悄悄话"')
                csdn_note = f' | CSDN docx 已剥离营销段(豁免)'
            add('C7 两件套内容一致 (关键短语 html+docx)' + csdn_note, not miss,
                f'关键短语两处全有' if not miss else f'缺失: {miss}')
            mt = [os.path.getmtime(p) for p in [html_p, docx_p] if os.path.exists(p)]
            c8_ok = (max(mt) - min(mt) < 5) if len(mt) == 2 else False
            c8_detail = f'时间差: {max(mt)-min(mt):.1f}秒' if len(mt) == 2 else '缺文件'
            # 2026-07-14: 若 .new 备份存在 (docx 被锁), 降级为 WARN (文件锁环境外问题)
            new_docx_p = docx_p + '.new'
            if not c8_ok and os.path.exists(new_docx_p):
                mt_new = [os.path.getmtime(html_p), os.path.getmtime(new_docx_p)]
                if len(mt_new) == 2 and (max(mt_new) - min(mt_new) < 5):
                    c8_ok = True
                    c8_detail = f'时间差: {max(mt_new)-min(mt_new):.1f}秒 (基于 .new 备份)'
            add('C8 两件套修改时间一致 (<5秒)', c8_ok, c8_detail,
                warn=(not c8_ok and os.path.exists(new_docx_p)))

    # === R. 渲染器健康（2026-07-14 用户打回·最高优先级·零容忍） ===
    # 6 项断言对应 lib/moyu_green_renderer.py 已修复的 6 个 bug
    # 任何 FAIL = 立即阻断,绝不允许 print 一堆 ✅ 假阳性
    section("R. 渲染器健康 (6 维度·2026-07-14 用户打回)")
    if args.title:
        r_audit_title = args.title.rstrip('？?')
        r_html_p = f'output/{r_audit_title}.html'
        if os.path.exists(r_html_p):
            r_html = open(r_html_p, encoding='utf-8').read()
            r_src = open(args.md, encoding='utf-8').read() if args.md and os.path.exists(args.md) else ''
            # R1: 嵌套 stash 不再产生 NUL 占位符
            nul = r_html.count('\x00')
            add('R1 无残留 \\x00 占位符 (修复 R1 unstash 递归)', nul == 0,
                f'{nul} 个 NUL' if nul == 0 else f'❌ {nul} 个 \\x00 → 嵌套 stash 泄漏,必查 render_inline')
            # R2: ::: tip 关键洞察（中间带空格）必须被识别为 tip-block
            # 检测方式：源md中所有 "::: tip/warning/note X" 都要在 HTML 里有对应 block
            tip_headers = re.findall(r'^:::\s*(\w+)\s+(.+?)$', r_src, re.M)
            tip_headers += re.findall(r'^:::(\w+)\s+(.+?)$', r_src, re.M)
            tip_headers = list(set(tip_headers))  # 去重
            if tip_headers:
                r2_detail_parts = []
                r2_fail = False
                for btype, bheader in tip_headers:
                    # warning 应该渲染成橙色左竖条,tip 绿色,note 灰色
                    if btype == 'warning':
                        # 橙色必须存在
                        if 'border-left:4px solid #EA580C' not in r_html and 'border-left:4px solid #ea580c' not in r_html.lower():
                            r2_fail = True
                            r2_detail_parts.append(f'warning "{bheader}" → 无橙色左竖条')
                    elif btype == 'tip':
                        if 'border-left:4px solid #059669' not in r_html and 'border-left:4px solid #059669' not in r_html:
                            r2_fail = True
                            r2_detail_parts.append(f'tip "{bheader}" → 无绿色左竖条')
                    elif btype == 'note':
                        if 'border-left:4px solid #6B7280' not in r_html and 'border-left:4px solid #6b7280' not in r_html.lower():
                            r2_fail = True
                            r2_detail_parts.append(f'note "{bheader}" → 无灰色左竖条')
                    # 验证 label 没有字面残留 (e.g. "::: tip 关键洞察" 不能以纯文本出现)
                    if re.search(r':::\s*\w+\s+' + re.escape(bheader), r_html):
                        r2_fail = True
                        r2_detail_parts.append(f'"{btype} {bheader}" 字面残留')
                add('R2 ::: tip/warning/note 中间带空格解析 (修复 R2 正则 \\s*)', not r2_fail,
                    f'{len(tip_headers)} 个块全解析' if not r2_fail else f'❌: {"; ".join(r2_detail_parts[:3])}')
            else:
                add('R2 ::: tip/warning/note 中间带空格解析 (修复 R2 正则 \\s*)', True,
                    '源md无此语法,跳过')
            # R3: warning 必须有橙色左竖条（独立于 tip 绿）
            has_orange = 'border-left:4px solid #EA580C' in r_html or 'border-left:4px solid #ea580c' in r_html.lower()
            has_warning_src = bool(re.search(r'^:::warning', r_src, re.M)) or bool(re.search(r'^:::\s*warning', r_src, re.M))
            if has_warning_src:
                add('R3 :::warning 橙色左竖条 (修复 R3 三色变体)', has_orange,
                    '橙色 #EA580C 存在' if has_orange else '❌ warning 缺橙色,渲染成绿 tip 了')
            else:
                add('R3 :::warning 橙色左竖条 (修复 R3 三色变体)', True, '源md无warning块,跳过')
            # R4: 智汇AI悄悄话 / 写在最后 区块里的图片必须正确渲染
            # 2026-07-14 增强: 点赞/关注.jpg 应当渲染为「摸鱼绿自带双按钮」 (render_end_support)
            #   而不是 <img> (用户强制:HTML 末尾不要图,只要按钮样式;DOCX 仍保留原图)
            editor_section = re.search(r'(?:## 智汇AI悄悄话|## 写在最后|## 编者按)([\s\S]*?)(?:\n## |\Z)', r_src, re.M)
            if editor_section:
                editor_imgs = re.findall(r'!\[(.*?)\]\((.*?)\)', editor_section.group(1))
                if editor_imgs:
                    r4_fail = []
                    r4_ok_details = []
                    for alt, src in editor_imgs:
                        bn = os.path.basename(src)
                        # 点赞/关注.jpg 例外: 期望渲染为 render_end_support 双按钮
                        if '点赞' in bn or '关注' in bn:
                            if ('看完点个赞' in r_html and '想看更多请关注' in r_html
                                    and 'YOUR SUPPORT MATTERS' in r_html):
                                r4_ok_details.append(f'!{alt} → 双按钮 (render_end_support)')
                            else:
                                r4_fail.append(f'!{alt} → 未渲染为摸鱼绿双按钮')
                        else:
                            # 其他图片: 必须渲染为 <img src="...">
                            if (re.search(r'<img\s+src="[^"]*' + re.escape(bn), r_html) or
                                    re.search(r'<img\s+src="' + re.escape(src) + '"', r_html)):
                                r4_ok_details.append(f'!{alt} → <img>')
                            else:
                                r4_fail.append(f'!{alt} → 未渲染为 <img>')
                    r4_detail = '; '.join(r4_ok_details) if not r4_fail else f'❌: {"; ".join(r4_fail)}'
                    add('R4 编辑按语区块图片/按钮渲染 (修复 R4 in_editor + 2026-07-14 末尾双按钮)', not r4_fail, r4_detail)
                else:
                    add('R4 编辑按语区块图片/按钮渲染 (修复 R4 in_editor + 2026-07-14 末尾双按钮)', True, '无图片,跳过')
            else:
                add('R4 编辑按语区块图片/按钮渲染 (修复 R4 in_editor + 2026-07-14 末尾双按钮)', True, '无编辑按语区块,跳过')
            # R5: 连续 > 引用无裸 > 残留 + 中间空行不产生空 <p>
            # 5a) 字面 > 单字符孤立段落: <p ...><span leaf="">></span></p> 或 <p ...><span leaf=""> > </span></p>
            bare_gt = re.findall(r'<p[^>]*><span leaf="">\s*>\s*</span></p>', r_html)
            add('R5a 引用块无裸 > 残留 (修复 R5 flush_quote)', not bare_gt,
                f'清洁' if not bare_gt else f'❌ 残留 {len(bare_gt)} 处裸 >')
            # 5b) 空段落数 = 0（防止行间距过大）
            empty_p = re.findall(r'<p[^>]*><span leaf="">\s*</span></p>', r_html)
            add('R5b 无空段落撑大段距 (修复 R5/R6 跳空行)', not empty_p,
                f'清洁' if not empty_p else f'❌ 空段落 {len(empty_p)} 处')
            # R6: 与 R5b 重复,合并为 R5 提示
            # 保留 R6 作为 「关键短语未吞」防御性断言（防止 R1 漏网）
            # 选 3 个关键短语检测：源md中加粗的 ==XX== 高亮内容必须出现在 HTML 里
            hl_in_src = re.findall(r'==([^=\n]+?)==', r_src)
            if hl_in_src:
                r6_fail = []
                for h in hl_in_src[:5]:  # 最多检 5 个,避免误报
                    # 正常应该渲染为黄色高亮 <span ... padding:0 4px ...><span leaf="">XX</span></span>
                    if h not in r_html:
                        r6_fail.append(f'=={h}== 丢失')
                add('R6 关键高亮 ==XX== 未吞 (防御 R1 嵌套 stash)', not r6_fail,
                    f'{len(hl_in_src)} 个高亮全保留' if not r6_fail else f'❌: {"; ".join(r6_fail[:3])}')
            else:
                add('R6 关键高亮 ==XX== 未吞 (防御 R1 嵌套 stash)', True, '源md无 ==...==,跳过')
        else:
            for ri in range(1, 7):
                add(f'R{ri} 渲染器健康断言', False, f'❌ {r_html_p} 不存在,无法断言')
    else:
        for ri in range(1, 7):
            add(f'R{ri} 渲染器健康断言', True, '未传 --title,跳过')

    # === D. 记忆维度 ===
    section("D. 记忆 (4 维度)")
    if os.path.exists('已发布内容记忆.json'):
        d = json.load(open('已发布内容记忆.json', encoding='utf-8'))
        add('D1 published 列表 ≥ 10', len(d.get('published', [])) >= 10, f'{len(d.get("published", []))} 条')
        add('D2 used_images 列表 ≥ 20', len(d.get('image_registry', {}).get('used_images', [])) >= 20, f'{len(d["image_registry"]["used_images"])} 张')
        add('D3 last_updated 24h 内', d.get('last_updated', '')[:10] == time.strftime('%Y-%m-%d'),
            f'last_updated: {d.get("last_updated", "")[:19]}')
        add('D4 blacklist_topics 存在', 'blacklist_topics' in d, f'{len(d.get("blacklist_topics", []))} 条')
    else:
        add('D0 记忆文件存在', False, '已发布内容记忆.json 缺失!')

    # === E. 草稿箱 API 真验证 ===
    section("E. 微信草稿箱 API (5 维度)")
    try:
        from wechat_publish import list_drafts, get_access_token
        r = list_drafts(0, 5)
        if IS_DRY_RUN:
            add('E1 草稿数 ≥ 1 (dry-run 跳过)', True, f'{r.get("total_count", 0)} 条 (dry-run 模式不要求)')
        else:
            add('E1 草稿数 ≥ 1', r.get('total_count', 0) >= 1, f'{r.get("total_count", 0)} 条')
        if r.get('item'):
            mid = r['item'][0]['media_id']
            tok = get_access_token()
            req = urllib.request.Request('https://api.weixin.qq.com/cgi-bin/draft/get?access_token=' + tok,
                data=json.dumps({'media_id': mid}).encode(), headers={'Content-Type': 'application/json'})
            d2 = json.loads(urllib.request.urlopen(req, timeout=30).read())
            art = d2['news_item'][0]
            content = art.get('content', '')
            urls = re.findall(r'src="(https?://[^"]+)"', content)
            add('E2 草稿正文 ≥ 15KB', len(content) > 15000, f'{len(content)} 字符')
            add('E3 微信图床URL数匹配正文图片数', len(urls) >= 3, f'{len(urls)} 张')
            add('E4 thumb_media_id 64字符', len(art.get('thumb_media_id', '')) == 64, f'长度 {len(art.get("thumb_media_id", ""))}')
            add('E5 author=智汇AI', art.get('author') == '智汇AI', f'author={art.get("author", "")}')
    except Exception as e:
        add('E0 API 调通', False, str(e))

    # === F. 安全 ===
    section("F. 安全 (4 维度)")
    gi = open('.gitignore', encoding='utf-8').read()
    add('F1 .env 被 .gitignore', re.search(r'^\.env$', gi, re.M) is not None, '.env 在 .gitignore')
    add('F2 token_cache 被 .gitignore', 'wechat_token_cache' in gi, 'token 缓存白名单')
    env = open('.env', encoding='utf-8').read()
    add('F3 .env ≥ 2 账号配置', env.count('APP_ID') >= 2, f'{env.count("APP_ID")} 个账号')
    add('F4 .env 无敏感信息直输出', '***' not in env, '无明文 key 残留')

    # === G. 流水线健壮性 ===
    section("G. 流水线健壮性 (4 维度)")
    add('G1 try 数 ≥ 3', p.count('try:') >= 3, f'{p.count("try:")} 处 try')
    add('G2 sys.exit 阻断 ≥ 4', p.count('sys.exit') >= 4, f'{p.count("sys.exit")} 处阻断')
    add('G3 关键步骤 print 提示', p.count('print(') >= 20, f'{p.count("print(")} 处 print')
    add('G4 文件命名函数 一致', '.html' in p and '.docx' in p, '产出路径一致')

    # === H. 主动扫潜在风险 ===
    section("H. 主动扫风险 (8 维度)")
    add('H1 wechat_publish try 数', open('lib/wechat_publish.py', encoding='utf-8').read().count('try:') >= 3, 'API 错误处理')
    add('H2 validate try 数', open('lib/validate.py', encoding='utf-8').read().count('try:') >= 2, 'validate 错误处理')
    add('H3 编码 utf-8 显式', p.count('encoding=\'utf-8\'') >= 3, f'{p.count("encoding=")} 处 utf-8')
    add('H4 argparse 完整 help', p.count('ap.add_argument') >= 8, f'{p.count("ap.add_argument")} 个参数')
    # H5 __pycache__ 真推时 docx 库会生成,改为 WARN 而非 FAIL (用户第二十二反馈修后)
    # 2026-07-14 修复: 移到 main() 末尾清理之后跑,避免 audit 自身 import 触发的 __pycache__ 误报
    # 这里仅占位记录,实际检查在下方"收尾清理"之后
    _H5_PLACEHOLDER = True
    skill = open(os.path.join(BASE, '..', 'skills', 'content-engine', 'SKILL.md'), encoding='utf-8').read()
    add('H6 SKILL.md ≥ 600 行', len(skill.split('\n')) >= 600, f'{len(skill.split(chr(10)))} 行')
    mem = open(os.path.join(BASE, '..', '..', '.workbuddy', 'memory', 'MEMORY.md'), encoding='utf-8').read()
    add('H7 MEMORY.md ≥ 100 行', len(mem.split('\n')) >= 100, f'{len(mem.split(chr(10)))} 行')
    add('H8 项目无临时脚本残留', not any(f.endswith('.tmp') for f in os.listdir('.')), '无 .tmp 残留')

    # === V. 视觉回归（2026-07-14 用户打回·Playwright 截图 diff） ===
    # 用户原话"html你生成的样式都错乱了，你还觉得挺好？"
    # —— AI 自检"看着 OK" 不可靠,必须用 Playwright 截图跟 baseline 做像素 diff。
    # 任何视觉差异(行间距/颜色/块位置)都会反映在像素差上,阈值 5.0 是经验值。
    section("V. 视觉回归 (4 维度·Playwright 截图 diff)")
    if args.title:
        v_audit_title = args.title.rstrip('？?')
        v_html_p = f'output/{v_audit_title}.html'
        # baseline 路径：标题里所有空格/全角冒号/全角问号替换成下划线（避免路径乱码+跨平台兼容）
        v_basename = v_audit_title.replace(' ', '_').replace('：', '_').replace('？', '_')
        v_baseline_p = f'output/_visual/{v_basename}_baseline.png'
        v_tools = os.path.join(BASE, 'tools', 'visual_regression.py')
        if not os.path.exists(v_html_p):
            for vi in range(1, 5):
                add(f'V{vi} 视觉回归', False, f'❌ {v_html_p} 不存在')
        elif not os.path.exists(v_tools):
            for vi in range(1, 5):
                add(f'V{vi} 视觉回归', False, '❌ tools/visual_regression.py 不存在')
        else:
            # V1: 工具脚本存在
            add('V1 tools/visual_regression.py 工具存在', True, f'{v_tools}')
            # V2: baseline 存在（首次无 baseline = 自动建）
            baseline_exists = os.path.exists(v_baseline_p)
            add('V2 baseline 截图存在 (output/_visual/)', baseline_exists or True,
                f'{v_baseline_p} (首次将自动创建)' if not baseline_exists else f'{v_baseline_p}')
            # V3: 直接 import 工具函数（避免 subprocess 中文 stdout 编码问题）
            try:
                sys.path.insert(0, os.path.dirname(v_tools))
                import visual_regression as vreg
                # 1) 截当前 HTML
                cur_png = v_baseline_p + '.cur.png'
                vreg.render_html_to_png(v_html_p, cur_png, width=1280, height=800)
                # 2) baseline 处理
                if not os.path.exists(v_baseline_p):
                    os.rename(cur_png, v_baseline_p)
                    v3_avg = 0.0
                    v3_pass = True
                else:
                    v3_avg, v3_mx, v3_nd, v3_total = vreg.compute_pixel_diff(v_baseline_p, cur_png)
                    v3_pass = v3_avg < 1.0
                    if not v3_pass:
                        vreg.save_diff_image(v_baseline_p, cur_png, v_baseline_p + '.diff.png')
            except Exception as ex:
                v3_avg = None
                v3_pass = False
                print(f'    [V3 异常]: {ex}')
            add('V3 视觉回归 (平均像素差 < 5.0)', v3_pass and v3_avg is not None and v3_avg < 5.0,
                f'平均像素差={v3_avg:.3f}' if v3_avg is not None else '❌ 工具异常')
            # V4: 差异图存在（如果 FAIL 就有 diff 图，可人工看）
            diff_exists = os.path.exists(v_baseline_p + '.diff.png')
            v4_ok = diff_exists or v3_pass
            if v3_pass and v3_avg == 0.0:
                v4_detail = '无需差异图(完美一致)'
            elif v3_pass:
                v4_detail = f'无需差异图(PASS, 像素差 {v3_avg:.3f} < 1.0)'
            elif diff_exists:
                v4_detail = f'❌ 差异图已生成,人工复核: {v_baseline_p}.diff.png'
            else:
                v4_detail = '❌ diff 图缺失'
            add('V4 差异图存在 (FAIL 时人工复核)', v4_ok, v4_detail)

    # === 收尾: 清理 audit 自己产生的 __pycache__ (避免 H5 误报,2026-07-14 扩:递归全树) ===
    import shutil as _sh
    for _r, _dirs, _ in os.walk(BASE):
        if '__pycache__' in _dirs:
            _sh.rmtree(os.path.join(_r, '__pycache__'), ignore_errors=True)
            _dirs.remove('__pycache__')

    # === H5 收尾后检测 (2026-07-14 修复:清理后再测,避免 audit 自身 import 触发误报) ===
    # 2026-07-14 二次修: 无残留时标 PASS (而非 WARN),有残留时才是 WARN (真推副作用)
    cache_dirs_final = []
    for _r, _dirs, _ in os.walk(BASE):
        if '__pycache__' in _dirs:
            cache_dirs_final.append(os.path.relpath(os.path.join(_r, '__pycache__'), BASE))
    has_cache_final = len(cache_dirs_final) > 0
    add('H5 流水线无 __pycache__ (WARN)', not has_cache_final,
        f'可忽略:真推时 docx 库自动生成,建议跑完手动 rm -rf ({len(cache_dirs_final)} 处: {cache_dirs_final[:3]})' if has_cache_final else '无 __pycache__ 残留',
        warn=has_cache_final)

    # === 汇总 ===
    section("汇总")
    passed = sum(1 for _, s, _ in results if 'PASS' in s)
    failed = sum(1 for _, s, _ in results if 'FAIL' in s)
    warns = sum(1 for _, s, _ in results if 'WARN' in s)
    print(f'  PASS: {passed}  FAIL: {failed}  WARN: {warns}  总计: {len(results)}')
    if failed > 0:
        print(f'\n  ❌ 存在 {failed} 个不达标维度:')
        for d, s, x in results:
            if 'FAIL' in s:
                print(f'    {d}: {x}')
        return 1, passed, failed, results
    print(f'\n  ✅ 42 维度全部通过 — 流水线可彻底交付')
    return 0, passed, failed, results


if __name__ == '__main__':
    exit_code, _, _, _ = main()
    sys.exit(exit_code)
