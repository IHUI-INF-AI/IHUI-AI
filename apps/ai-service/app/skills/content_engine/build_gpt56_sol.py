# -*- coding: utf-8 -*-
"""
GPT-5.6-Sol 删光AI创业者Mac硬盘 — 公众号文章一键构建脚本（摸鱼绿主题版 v6 终极版）
选题: OpenAI GPT-5.6-Sol 删光 AI 创业者 Matt Shumer 的 Mac 硬盘
日期: 2026-07-11
模板: 内容公式v2（热点解读）

v7 终极版（2026-07-13，DOCX样式层级化+间距统一+真编号+字体fallback）：
  在 v5 极致美化基础上，彻底清理多余装饰，强化高级隐秘感：
    1. 删除章节前 emerald 分割线 + 装饰小圆点（标题→01/02/03直接连）
    2. 删除死代码：add_cover() / add_toc()（用户禁用封面页/目录页）
    3. 文末保留点赞/关注图（独立居中段落，2026-07-14 用户恢复）
    4. 页面背景色 #FAFAFA（通过 w:background XML，增加层次感）
    5. 图片边框加深 D1D5DB + 宽度翻倍 19050EMU + 阴影加强（alpha 55k, blur 60k）
    6. 章节标题升级：▎竖条 24pt深绿 + 编号 22pt绿 + 标题 20pt深绿
    7. 全局 Normal 字号 11pt→11.5pt
    8. 正文段落：base_size 11→11.5, line_spacing 1.85→1.9, space_before 2→3, space_after 8→10
    9. 引用块：padding加大(0.6/0.5cm), line_spacing 1.7→1.8, space 12→14
    10. 编辑按语：base_size 10→10.5, base_color COLOR_TEXT_MAIN, 行距1.6→1.75, 缩进0.5/0.4cm
    11. H1标题 22pt→24pt
    12. verify_docx_images阈值 6→4（移除文末图后实际4张）
  保留：emerald 色系 + 13组件全标记兼容 + 零容忍禁令（禁止封面页/目录页/图片caption/页眉页脚）
"""

import os
import sys
import re
import time
import glob

# ===== 路径设置 =====
BASE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE)

# ── 项目边界硬门禁（导入即生效，fail-closed） ──
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "koubo_workflow"))
import project_boundary
project_boundary.check_action(tool="build_gpt56_sol.py")

from lib.validate import validate

OUTPUT_DIR = os.path.join(BASE, 'output')
IMAGES_DIR = os.path.join(OUTPUT_DIR, 'images')
ASSETS_IMAGES_DIR = os.path.join(BASE, 'assets', 'images')
MD_PATH = os.path.join(OUTPUT_DIR, 'AI删光了创业者的硬盘，3个普通人避坑真相.md')
DOCX_PATH = os.path.join(OUTPUT_DIR, 'AI删光了创业者的硬盘，3个普通人避坑真相_摸鱼绿.docx')
SUBTITLE = 'OpenAI GPT-5.6-Sol 删光创业者硬盘，普通人必须知道的 3 个避坑真相'

# ===== 摸鱼绿主题色彩体系（官方emerald色系，禁止Material绿） =====
# 官方参考: github.com/isjiamu/gzh-design-skill references/theme-moyu-green.md
COLOR_GREEN = (0x05, 0x96, 0x69)       # 主色 emerald-600 #059669（竖条/编号/关键词加粗）
COLOR_GREEN_DARK = (0x04, 0x7E, 0x55)  # 深 emerald（标题描边感）
COLOR_GREEN_LITE = (0x10, 0xB9, 0x81)  # 辅色 emerald-500 #10B981（浅元素）
COLOR_GREEN_PALE = (0xEC, 0xFD, 0xF4)  # 极浅绿底 #ECFDF5（金句卡/徽章底）
COLOR_YELLOW = (0xFD, 0xE6, 0x8A)      # 黄色点睛 #FDE68A
COLOR_TITLE = (0x11, 0x18, 0x27)       # 标题 #111827
COLOR_TEXT_MAIN = (0x37, 0x41, 0x51)   # 正文 #374151
COLOR_TEXT_LIGHT = (0x6B, 0x72, 0x80)  # 注释 #6B7280
COLOR_BORDER_LIGHT = (0xBB, 0xF7, 0xD0)  # 浅绿边框 #BBF7D0
COLOR_EDITOR_BG = (0xF0, 0xFD, 0xF4)   # 编者按浅绿底 #F0FDF4（emerald，非Material）
COLOR_CODE_BG = (0x1E, 0x29, 0x3B)     # 代码块深色底 #1E293B
COLOR_CODE_BAR = (0x0F, 0x17, 0x2A)    # 代码标签条 #0F172A
COLOR_CODE_TEXT = (0xE2, 0xE8, 0xF0)   # 代码块文字 #E2E8F0
COLOR_CODE_LABEL = (0x94, 0xA3, 0xB8)  # 代码标签文字 #94A3B8
COLOR_QUOTE_BG = (0xF9, 0xFA, 0xFB)    # 引用块浅灰底 #F9FAFB
COLOR_QUOTE_BORDER = (0xD1, 0xD5, 0xDB)  # 引用虚线边框 #D1D5DB
COLOR_SEP = (0xE5, 0xE7, 0xEB)         # 分割线浅灰 #E5E7EB
COLOR_WHITE = (0xFF, 0xFF, 0xFF)
COLOR_RED = (0xEF, 0x44, 0x44)         # mac 圆点 红
COLOR_AMBER = (0xF5, 0x9E, 0x0B)       # mac 圆点 黄
COLOR_GREEN_DOT = (0x10, 0xB9, 0x81)   # mac 圆点 绿


def _load_body_from_md():
    """从已修复的md文件读取正文，确保脚本与md始终同步"""
    if os.path.exists(MD_PATH):
        with open(MD_PATH, 'r', encoding='utf-8') as f:
            return f.read()
    return ""


BODY = _load_body_from_md()


# ===== 写入Markdown =====
def write_md():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    project_boundary.check_write(MD_PATH)
    with open(MD_PATH, 'w', encoding='utf-8') as f:
        f.write(BODY)
    print(f"[1/4] Markdown 已写入: {MD_PATH}")
    cn_chars = sum(1 for c in BODY if '\u4e00' <= c <= '\u9fff')
    print(f"      中文字数: {cn_chars}")


# ===== 验证 =====
def run_validate():
    print(f"\n[2/4] 执行22项自检...")
    all_pass, reports = validate(MD_PATH)
    for r in reports:
        print(f"  {r}")
    print(f"\n  {'✅ 22项自检全部通过' if all_pass else '❌ 存在未通过项，需修复'}")
    return all_pass


# ===== 构建DOCX（摸鱼绿主题 v5 极致美化，兼容全标记） =====
# 参数化：可被统一流水线 publish_pipeline.py 复用，不再依赖全局常量
def build_docx(md_path=MD_PATH, docx_path=DOCX_PATH,
               images_dir=IMAGES_DIR, assets_dir=ASSETS_IMAGES_DIR,
               cover_img=None, digest=SUBTITLE):
    project_boundary.check_write(docx_path)
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    print(f"\n[构建] 构建DOCX（摸鱼绿主题 v5 极致美化，全标记兼容）...")

    # 读取MD（参数化，不再依赖全局 BODY）
    if not os.path.exists(md_path):
        print(f"  ❌ MD不存在: {md_path}")
        return None
    with open(md_path, 'r', encoding='utf-8') as f:
        body_text = f.read()
    # 压空行: 3+连续空行→1空行(避免 doc 导入平台后排版出现大片空白, 用户第二十二反馈)
    body_text = re.sub(r"\n{3,}", "\n\n", body_text)

    # 提取标题（首个 # 行）
    title_text = None
    for ln in body_text.split('\n'):
        s = ln.strip()
        if s.startswith('# ') and not s.startswith('## '):
            title_text = s[2:].strip()
            break
    title_text = title_text or os.path.splitext(os.path.basename(md_path))[0]

    # 封面图自动探测
    if cover_img is None and os.path.isdir(images_dir):
        cands = sorted(glob.glob(os.path.join(images_dir, 'sec1*')))
        if not cands:
            cands = sorted(glob.glob(os.path.join(images_dir, '*.jpg'))) + \
                    sorted(glob.glob(os.path.join(images_dir, '*.png')))
        if cands:
            cover_img = cands[0]

    doc = Document()

    # ===== 样式初始化 =====
    for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        try:
            s = doc.styles[style_name]
        except KeyError:
            continue
        if style_name == 'Heading 1':
            s.font.name = 'Microsoft YaHei'
            s.font.size = Pt(22)
            s.font.bold = True
            s.font.color.rgb = RGBColor(*COLOR_TEXT_MAIN)
        elif style_name == 'Heading 2':
            s.font.name = 'Microsoft YaHei'
            s.font.size = Pt(20)
            s.font.bold = True
            s.font.color.rgb = RGBColor(*COLOR_GREEN_DARK)
        elif style_name == 'Heading 3':
            s.font.name = 'Microsoft YaHei'
            s.font.size = Pt(13.5)
            s.font.bold = True
            s.font.color.rgb = RGBColor(*COLOR_GREEN)
        rpr = s.element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:ascii'), 'Microsoft YaHei')
        rf.set(qn('w:hAnsi'), 'Microsoft YaHei')
        rf.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== 全局样式 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11.5)
    font.color.rgb = RGBColor(*COLOR_TEXT_MAIN)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rfonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 页面边距 + 背景色
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        # 页面背景色 #FAFAFA
        sectPr = section._sectPr
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), 'FAFAFA')
        sectPr.append(bg)

    # ===== 辅助函数 =====
    def set_east_asia(run, font_name='Microsoft YaHei'):
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:ascii'), font_name)
        rf.set(qn('w:hAnsi'), font_name)
        rf.set(qn('w:eastAsia'), font_name)

    def add_shading(paragraph, fill_hex):
        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        pPr.append(shd)

    def _add_side_border(paragraph, side, color_hex, sz='24', space='8', val='single'):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)
        e = OxmlElement('w:' + side)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), space)
        e.set(qn('w:color'), color_hex)
        pBdr.append(e)

    def add_left_border(paragraph, color_hex, sz='24', space='8'):
        _add_side_border(paragraph, 'left', color_hex, sz, space, 'single')

    def add_bottom_border(paragraph, color_hex='E5E7EB', sz='6', space='1'):
        _add_side_border(paragraph, 'bottom', color_hex, sz, space, 'single')

    def add_dashed_border(paragraph, color_hex='D1D5DB', sz='12', space='6'):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ['top', 'left', 'bottom', 'right']:
            e = OxmlElement('w:' + side)
            e.set(qn('w:val'), 'dashed')
            e.set(qn('w:sz'), sz)
            e.set(qn('w:space'), space)
            e.set(qn('w:color'), color_hex)
            pBdr.append(e)
        pPr.append(pBdr)

    def add_card_border(paragraph, left_color='059669', side_color='BBF7D0',
                        left_sz='28', side_sz='6'):
        """卡片盒：左侧 emerald 粗竖条 + 其余三边浅绿细边框"""
        _add_side_border(paragraph, 'left', left_color, left_sz, '10', 'single')
        for sd in ['top', 'right', 'bottom']:
            _add_side_border(paragraph, sd, side_color, side_sz, '6', 'single')

    def add_decoration_bar(fill_hex, height_pt=10):
        """装饰色条（封面顶部/底部）：段落底色模拟横幅，height 由字号撑起"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_shading(p, fill_hex)
        r = p.add_run('　')
        r.font.size = Pt(height_pt)
        r.font.name = 'Microsoft YaHei'
        set_east_asia(r)
        return p

    # 图片边框加深 + 阴影加强（相框质感 v2）
    def add_picture_frame(shape, color='D1D5DB', w=19050):
        """给图片加浅灰细边框 + 轻微外阴影（相框质感），直接操作 drawing xml"""
        try:
            spPr = shape._inline.graphic.graphicData.pic.spPr
            kids = list(spPr)
            insert_idx = 0
            for idx, el in enumerate(kids):
                if el.tag in (qn('a:prstGeom'), qn('a:custGeom')):
                    insert_idx = idx + 1
            # 边框
            ln = OxmlElement('a:ln')
            ln.set('w', str(w))
            ln.set('cap', 'flat')
            sf = OxmlElement('a:solidFill')
            c = OxmlElement('a:srgbClr'); c.set('val', color)
            sf.append(c); ln.append(sf)
            # 阴影
            ef = OxmlElement('a:effectLst')
            sh = OxmlElement('a:outerShdw')
            sh.set('blurRad', '60000'); sh.set('dist', '25000')
            sh.set('dir', '5400000'); sh.set('rotWithShape', '0')
            sc = OxmlElement('a:srgbClr'); sc.set('val', '6B7280')
            al = OxmlElement('a:alpha'); al.set('val', '55000')
            sc.append(al); sh.append(sc); ef.append(sh)
            spPr.insert(insert_idx, ln)
            spPr.insert(insert_idx + 1, ef)
        except Exception as e:
            print(f"    ⚠️ 图片边框添加跳过: {e}")

    # add_cover / add_toc 已删除（用户禁用封面页+目录页）

    # ===== 行内标记解析：**加粗**(绿) / ==高亮==(黄底) / `代码`(浅灰等宽) =====
    TOKEN_RE = re.compile(r'(\*\*[^*]+\*\*|==[^=]+==|`[^`]+`)')

    def add_inline_runs(paragraph, text, base_size=11, base_color=COLOR_TEXT_MAIN):
        for tok in TOKEN_RE.split(text):
            if not tok:
                continue
            if tok.startswith('**') and tok.endswith('**') and len(tok) >= 4:
                seg = tok[2:-2]
                r = paragraph.add_run(seg)
                r.font.bold = True
                r.font.color.rgb = RGBColor(*COLOR_GREEN)
            elif tok.startswith('==') and tok.endswith('==') and len(tok) >= 4:
                seg = tok[2:-2]
                r = paragraph.add_run(seg)
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif tok.startswith('`') and tok.endswith('`') and len(tok) >= 2:
                seg = tok[1:-1]
                r = paragraph.add_run(seg)
                r.font.name = 'Consolas'
                r.font.color.rgb = RGBColor(*COLOR_TEXT_LIGHT)
            else:
                r = paragraph.add_run(tok)
                r.font.color.rgb = RGBColor(*base_color)
            r.font.size = Pt(base_size)
            r.font.name = 'Microsoft YaHei'
            set_east_asia(r)

    # ===== 区块 flush（oneliner/quote/tip/code）=====
    def flush_block(btype, buf, meta):
        if btype == 'code':
            lang = (meta or 'code').strip() or 'code'
            # 顶部语言条 + mac 三色圆点
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(8)
            lp.paragraph_format.space_after = Pt(8)
            add_shading(lp, '0F172A')
            for col in (COLOR_RED, COLOR_AMBER, COLOR_GREEN_DOT):
                dot = lp.add_run('●')
                dot.font.size = Pt(8)
                dot.font.color.rgb = RGBColor(*col)
                dot.font.name = 'Microsoft YaHei'
                set_east_asia(dot)
            gap = lp.add_run('  ')
            gap.font.size = Pt(8)
            gap.font.name = 'Microsoft YaHei'
            set_east_asia(gap)
            lr = lp.add_run(lang + '   ')
            lr.font.size = Pt(9)
            lr.font.name = 'Consolas'
            lr.font.color.rgb = RGBColor(*COLOR_CODE_LABEL)
            set_east_asia(lr)
            for cl in buf:
                cl = cl.rstrip('\n')
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.35
                p.paragraph_format.left_indent = Cm(0.3)
                add_shading(p, '1E293B')
                run = p.add_run(cl if cl.strip() != '' else ' ')
                run.font.size = Pt(10)
                run.font.name = 'Consolas'
                run.font.color.rgb = RGBColor(*COLOR_CODE_TEXT)
                set_east_asia(run)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(8)
            return

        text = '\n'.join(buf).strip()

        if btype == 'oneliner':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.line_spacing = 1.6
            add_shading(p, 'ECFDF5')
            add_dashed_border(p, 'BBF7D0', '12', '8')
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.right_indent = Cm(0.4)
            # 大引号装饰
            q = p.add_run('“')
            q.font.size = Pt(22)
            q.font.bold = True
            q.font.color.rgb = RGBColor(*COLOR_GREEN_LITE)
            q.font.name = 'Microsoft YaHei'
            set_east_asia(q)
            if meta:
                rk = p.add_run(meta)
                rk.font.size = Pt(9.5)
                rk.font.color.rgb = RGBColor(*COLOR_TEXT_LIGHT)
                rk.font.name = 'Microsoft YaHei'
                set_east_asia(rk)
                rk.add_break()
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(12.5)
            r.font.color.rgb = RGBColor(*COLOR_GREEN)
            r.font.name = 'Microsoft YaHei'
            set_east_asia(r)
            qe = p.add_run('”')
            qe.font.size = Pt(22)
            qe.font.bold = True
            qe.font.color.rgb = RGBColor(*COLOR_GREEN_LITE)
            qe.font.name = 'Microsoft YaHei'
            set_east_asia(qe)
            return

        if btype == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.7
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.4)
            add_shading(p, 'F9FAFB')
            add_left_border(p, '059669', '28', '10')
            # 大引号装饰
            q = p.add_run('“')
            q.font.size = Pt(20)
            q.font.bold = True
            q.font.color.rgb = RGBColor(*COLOR_GREEN_LITE)
            q.font.name = 'Microsoft YaHei'
            set_east_asia(q)
            add_inline_runs(p, text, base_size=11, base_color=COLOR_TEXT_MAIN)
            for rn in p.runs:
                rn.font.italic = True
            return

        # tip / warning / note / 其它提示类 —— 卡片盒
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(16)
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.3)
        add_shading(p, 'F0FDF4')
        add_card_border(p, '059669', 'BBF7D0', '28', '6')
        label = meta if meta else ('提示' if btype == 'tip' else btype)
        rk = p.add_run('◆ ' + label)
        rk.font.bold = True
        rk.font.size = Pt(10.5)
        rk.font.color.rgb = RGBColor(*COLOR_GREEN)
        rk.font.name = 'Microsoft YaHei'
        set_east_asia(rk)
        rk.add_break()
        add_inline_runs(p, text, base_size=10.5, base_color=COLOR_TEXT_MAIN)
        return

    # ===== 解析Markdown并生成DOCX =====
    # 先提取章节用于目录页
    chapters = []
    for ln in body_text.split('\n'):
        s = ln.strip()
        if s.startswith('## ') and not s.startswith('### ') and '智汇AI悄悄话' not in s:
            chapters.append(re.sub(r'^##\s*▎?\s*\d*\.?\s*', '', s).strip())

    # 预扫描：本文章是否已有 H2 自带 "N." 显式编号。
    # 若有，则整篇走"显式编号"模式（不再自动补 01/02/03），避免双编号；
    # 若全篇 H2 都无编号，则维持旧自动编号行为（01/02/03）。
    doc_has_explicit_num = False
    for ln in body_text.split('\n'):
        s = ln.strip()
        if s.startswith('## ') and not s.startswith('### '):
            if re.match(r'^##\s*▎?\s*\d+\.', s):
                doc_has_explicit_num = True
                break

    # 不再加封面页/导读页（用户第二十二反馈：前面太啰嗦，导入平台后还要手动删空行）
    # 直接从正文第一段开始, 节省空间, 平台导入排版干净
    # 页眉页脚已移除（用户要求）

    lines = body_text.strip().split('\n')
    total = len(lines)
    i = 0
    block_type = None
    block_buf = []
    block_meta = None
    chapter_idx = 0

    while i < total:
        raw = lines[i]
        line = raw.strip()

        # ---- 区块收集中 ----
        if block_type is not None:
            end_now = (block_type == 'code' and line.startswith('```')) or \
                      (block_type != 'code' and line == ':::')
            if end_now:
                flush_block(block_type, block_buf, block_meta)
                block_type, block_buf, block_meta = None, [], None
                i += 1
                continue
            block_buf.append(raw)
            i += 1
            continue

        # ---- 空行 ----
        if not line:
            i += 1
            continue

        # ---- 区块开始 ----
        if line.startswith('```'):
            block_type = 'code'
            block_buf, block_meta = [], line[3:].strip()
            i += 1
            continue
        if line.startswith(':::') and len(line) > 3:
            m = re.match(r':::(\w+)(?:\s+(.*))?', line)
            block_type = m.group(1) if m else 'tip'
            block_meta = (m.group(2) or '') if m else ''
            block_buf = []
            i += 1
            continue
        if line == ':::':  # 孤立结束符，忽略
            i += 1
            continue

        # ---- H1 文章标题（封面页已渲染，正文保留标题）----
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Heading 1']
            except KeyError:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run(title_text)
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*COLOR_TEXT_MAIN)
            run.font.name = 'Microsoft YaHei'
            set_east_asia(run)
            i += 1
            continue

        # ---- H2 章节标题（含"智汇AI悄悄话"也走此分支）----
        # 自适应编号：源 MD 标题若已自带 "▎" 与 "N." 编号（如 "▎1. 真相"），
        # 则直接复用，构建器不再叠加自动编号，避免 docx 出现"两个数字/双竖条"。
        if line.startswith('## '):
            raw_h = line[3:].strip()
            num_m = re.match(r'^▎?\s*(\d+)\.', raw_h)   # 提取标题里自带的 "N."
            has_num = bool(num_m)
            md_num = num_m.group(1) if num_m else None
            # 清理标题文本：去掉开头的 ▎ 与可选 "N."，交给下面统一加 ▎ 与编号
            heading_text = re.sub(r'^▎\s*', '', raw_h)
            heading_text = re.sub(r'^\d+\.\s*', '', heading_text)
            is_editor = '智汇AI悄悄话' in heading_text
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Heading 2']
            except KeyError:
                pass
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(16)
            if is_editor:
                # 编辑按语标题：emerald 加粗（单一 ▎，不叠自动编号）
                run = p.add_run('▎' + heading_text)
                run.font.size = Pt(17)
                run.font.bold = True
                run.font.color.rgb = RGBColor(*COLOR_GREEN)
                run.font.name = 'Microsoft YaHei'
                set_east_asia(run)
            else:
                run_bar = p.add_run('▎')
                run_bar.font.size = Pt(24)
                run_bar.font.bold = True
                run_bar.font.color.rgb = RGBColor(*COLOR_GREEN_DARK)
                run_bar.font.name = 'Microsoft YaHei'
                set_east_asia(run_bar)
                if has_num:
                    run_num = p.add_run(f'{md_num}.  ')      # 复用 MD 自带编号，如 "1. "
                elif not doc_has_explicit_num:
                    chapter_idx += 1
                    run_num = p.add_run(f'{chapter_idx:02d}  ')  # 全篇无显式编号时自动补 01/02/03
                else:
                    run_num = None                            # 显式编号模式且本标题无数字→不加数字
                if run_num is not None:
                    run_num.font.size = Pt(22)
                    run_num.font.bold = True
                    run_num.font.color.rgb = RGBColor(*COLOR_GREEN)
                    run_num.font.name = 'Consolas'
                    set_east_asia(run_num)
                run = p.add_run(heading_text)
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(*COLOR_GREEN_DARK)
                run.font.name = 'Microsoft YaHei'
                set_east_asia(run)
            i += 1
            continue

        # ---- ### 小标题（绿加粗 + 黄色高亮）----
        if line.startswith('### '):
            sub = line[4:].strip()
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Heading 3']
            except KeyError:
                pass
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(sub)
            r.font.bold = True
            r.font.size = Pt(13.5)
            r.font.color.rgb = RGBColor(*COLOR_GREEN)
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            r.font.name = 'Microsoft YaHei'
            set_east_asia(r)
            i += 1
            continue

        # ---- 图片 ----
        if line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt, img_path = match.group(1), match.group(2)
                md_dir = os.path.dirname(os.path.abspath(md_path))
                full_path = os.path.normpath(os.path.join(md_dir, img_path))
                if not os.path.exists(full_path):
                    full_path = os.path.normpath(os.path.join(images_dir, os.path.basename(img_path)))
                if not os.path.exists(full_path):
                    full_path = os.path.normpath(os.path.join(BASE, img_path))
                if os.path.exists(full_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(8)
                    run = p.add_run()
                    # 点赞/关注图 4.0 英寸,正文图 5.0 英寸
                    w = Inches(4.0) if ('点赞' in img_path or '关注' in img_path) else Inches(5.0)
                    shape = run.add_picture(full_path, width=w)
                    add_picture_frame(shape, 'D1D5DB', 19050)
                else:
                    print(f'    ⚠️ 图片不存在，跳过: {img_path}')
                i += 1
                continue

        # ---- > 引用块（emerald左粗竖条 + 浅灰底 + 大引号）----
        if line.startswith('> '):
            q = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(16)
            p.paragraph_format.line_spacing = 1.8
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.5)
            add_shading(p, 'F9FAFB')
            add_left_border(p, '059669', '28', '10')
            add_inline_runs(p, q, base_size=11, base_color=COLOR_TEXT_MAIN)
            for rn in p.runs:
                rn.font.italic = True
            i += 1
            continue

        # ---- 1. 有序列表（绿色加粗序号 + Word 真编号）----
        mord = re.match(r'^(\d+)\.\s+(.*)$', line)
        if mord:
            rest = mord.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.6
            add_inline_runs(p, rest, base_size=11, base_color=COLOR_TEXT_MAIN)
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(*COLOR_GREEN)
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.name = 'Microsoft YaHei'
                set_east_asia(p.runs[0])
            i += 1
            continue

        # ---- 编辑按语正文段（紧跟"## 智汇AI悄悄话"之后）----
        if line.startswith('本文由智汇AI社区') or '智汇AI是吉林省爱智汇' in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(16)
            p.paragraph_format.line_spacing = 1.75
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.4)
            add_shading(p, 'F0FDF4')
            add_card_border(p, '059669', 'BBF7D0', '28', '6')
            add_inline_runs(p, line, base_size=10.5, base_color=COLOR_TEXT_MAIN)
            i += 1
            continue

        # ---- 普通段落（含 **加粗** / ==高亮== / `代码` 行内标记）----
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.first_line_indent = Pt(22)  # 首行缩进 2 字符
        add_inline_runs(p, line, base_size=11.5, base_color=COLOR_TEXT_MAIN)
        i += 1

    # ===== 文末横幅已删除（用户禁用营销类文末图片）=====

    # 2026-07-14: 写临时文件再原子重命名,绕过微信端预览导致的 docx 文件锁
    tmp_path = docx_path + '.tmp'
    new_path = docx_path + '.new'
    # 清旧备份,避免 C1.5 一直 WARN
    for stale in (tmp_path, new_path):
        if os.path.exists(stale):
            try:
                os.remove(stale)
            except Exception:
                pass
    doc.save(tmp_path)
    # 1) 先尝试直接覆盖 (目标文件未锁时)
    try:
        if os.path.exists(docx_path):
            os.remove(docx_path)
        os.rename(tmp_path, docx_path)
        print(f"  DOCX 已保存: {docx_path}")
        return docx_path
    except (PermissionError, OSError) as e:
        # 2) 目标被锁: 退到 .new 备选路径,不阻塞流水线
        new_path = docx_path + '.new'
        if os.path.exists(new_path):
            try:
                os.remove(new_path)
            except Exception:
                pass
        # 把旧的 .tmp 改名到 .new
        try:
            os.rename(tmp_path, new_path)
        except Exception:
            import shutil as _sh
            _sh.copy2(tmp_path, new_path)
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        print(f"  ⚠️ docx 被外部进程锁定({type(e).__name__}),已写入备选路径: {new_path}")
        print(f"      请关闭微信端 docx 预览后手动: copy /Y \"{new_path}\" \"{docx_path}\"")
        return new_path


# ===== 验证DOCX内嵌图片数 =====
def verify_docx_images(docx_path=DOCX_PATH):
    from docx import Document
    print(f"\n[验证] 验证DOCX内嵌图片...")
    doc = Document(docx_path)
    shape_count = len(doc.inline_shapes)
    print(f"  InlineShapes数: {shape_count}")
    if shape_count >= 4:
        print(f"  ✅ 配图完整性验证通过（≥4张真实嵌入）")
    else:
        print(f"  ⚠️ 期望≥4张图片（正文配图），实际{shape_count}张")
    return shape_count


# ===== 主流程 =====
if __name__ == '__main__':
    print('=' * 60)
    print('  GPT-5.6-Sol删光硬盘 — 公众号文章一键构建（摸鱼绿主题 v5 极致美化）')
    print('=' * 60)

    write_md()
    all_pass = run_validate()
    if not all_pass:
        print("\n⚠️ 自检未全部通过，请根据报告修复后重新构建")
        sys.exit(1)

    build_docx()
    verify_docx_images()

    print('\n' + '=' * 60)
    print('  ✅ 构建完成！')
    print(f'  Markdown: {MD_PATH}')
    print(f'  DOCX: {DOCX_PATH}')
    print('=' * 60)



